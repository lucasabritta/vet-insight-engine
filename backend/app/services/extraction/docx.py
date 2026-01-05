from .base import DocumentExtractor, ExtractionResult
import docx

class DocxExtractor(DocumentExtractor):
    def extract(self, file_path: str) -> ExtractionResult:
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        meta = {"type": "docx", "paragraphs": len(doc.paragraphs)}
        return ExtractionResult(text=text, meta=meta)
